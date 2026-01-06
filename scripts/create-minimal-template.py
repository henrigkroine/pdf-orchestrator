#!/usr/bin/env python3
"""
Create minimal Word template to test Adobe merge functionality.
Uses absolute minimum fields to isolate the issue.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_merge_field_as_single_run(paragraph, field_text):
    """
    Add merge field ensuring it's in a single XML run.
    This is critical for Adobe to recognize it.
    """
    # Clear any existing runs
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    
    # Create new run element
    run = paragraph.add_run()
    
    # Set the text in one operation
    run.text = field_text
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    
    return run

# Create document
doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title - single field
title_para = doc.add_paragraph()
add_merge_field_as_single_run(title_para, '{{title}}')

# Blank line
doc.add_paragraph()

# Subtitle - single field
subtitle_para = doc.add_paragraph()
add_merge_field_as_single_run(subtitle_para, '{{subtitle}}')

# Save
output_path = 'T:/Projects/pdf-orchestrator/templates/word/teei-minimal-test.docx'
doc.save(output_path)
print(f'[OK] Minimal template created: {output_path}')
print('  Fields: {{title}}, {{subtitle}}')
