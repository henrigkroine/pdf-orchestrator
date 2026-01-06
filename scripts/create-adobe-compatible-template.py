#!/usr/bin/env python3
"""
Create Word template with Adobe Document Generation compatible merge fields.
Uses simple structure that matches our test data exactly.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Create new document
doc = Document()

# Set page margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title - exactly matching our data structure
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('{{title}}')
run.font.name = 'Arial'
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 57, 63)

# Add space
doc.add_paragraph()

# Subtitle - exactly matching our data structure
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('{{subtitle}}')
run.font.name = 'Arial'
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

# Add space
doc.add_paragraph()
doc.add_paragraph()

# Simple text content
content_para = doc.add_paragraph()
run = content_para.add_run('This template tests if python-docx templates can work with Adobe when fields match exactly.')
run.font.name = 'Arial'
run.font.size = Pt(11)

# Save template
output_path = 'T:/Projects/pdf-orchestrator/templates/word/teei-exact-match-test.docx'
doc.save(output_path)
print(f'[OK] Template created: {output_path}')
print('[INFO] Template has exact field match: title, subtitle')
print('[INFO] No array loops - just simple merge fields')
