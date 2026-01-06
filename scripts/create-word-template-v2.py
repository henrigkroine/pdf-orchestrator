#!/usr/bin/env python3
"""
Create a professional Word template for TEEI showcase documents.
Template uses Adobe Document Generation merge fields with simplified content iteration.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create new document
doc = Document()

# Set page margins (1 inch = 1 inch)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title with merge field
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('{{title}}')
run.font.name = 'Arial'
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 57, 63)  # TEEI Nordshore #00393F
title.paragraph_format.space_after = Pt(24)

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('{{subtitle}}')
run.font.name = 'Arial'
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(186, 143, 90)  # TEEI Gold #BA8F5A
subtitle.paragraph_format.space_after = Pt(36)

# Content section with loop - each array item is now accessed with correct context
# In Mustache, within {{#content}}, the current item is accessed with {{.}}
content_loop_start = doc.add_paragraph()
run = content_loop_start.add_run('{{#content}}')
run.font.name = 'Arial'
run.font.size = Pt(11)

# Content item - use {{.}} to reference current string in array
item = doc.add_paragraph()
run = item.add_run('{{.}}')
run.font.name = 'Arial'
run.font.size = Pt(11)
item.paragraph_format.space_after = Pt(12)

# Close content loop
close_loop = doc.add_paragraph()
run = close_loop.add_run('{{/content}}')
run.font.name = 'Arial'
run.font.size = Pt(11)

# Metadata section
doc.add_page_break()
metadata = doc.add_paragraph()
metadata.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = metadata.add_run('Author: {{metadata.author}}')
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

date_para = doc.add_paragraph()
run = date_para.add_run('Date: {{metadata.date}}')
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

org_para = doc.add_paragraph()
run = org_para.add_run('Organization: {{metadata.organization}}')
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

# Save template
output_path = 'T:/Projects/pdf-orchestrator/templates/word/teei-showcase-template.docx'
doc.save(output_path)
print(f'Template created: {output_path}')
